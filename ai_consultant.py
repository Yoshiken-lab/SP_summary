#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
スクールフォト売上分析システム - AIコンサルタント

ローカルLLM（Ollama）を使用して売上データを分析し、
経営アドバイスを生成する
出力形式: Word, PDF, Markdown
"""

import json
import requests
import subprocess
from pathlib import Path
from datetime import datetime
from database import get_connection
from analytics import get_all_analytics, get_current_fiscal_year

# オプション: Word出力用ライブラリ
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# 設定ファイルのパス
CONFIG_PATH = Path(__file__).parent / "config.json"


def load_config():
    """設定ファイルを読み込む"""
    if not CONFIG_PATH.exists():
        return None

    with open(CONFIG_PATH, 'r', encoding='utf-8') as f:
        return json.load(f)


def is_ollama_available(config):
    """Ollamaが利用可能かチェック"""
    if not config or not config.get('ai_consultant', {}).get('enabled'):
        return False

    ollama_config = config.get('ai_consultant', {}).get('ollama', {})
    base_url = ollama_config.get('base_url', 'http://localhost:11434')

    try:
        response = requests.get(f"{base_url}/api/tags", timeout=5)
        return response.status_code == 200
    except requests.exceptions.RequestException:
        return False


def get_summary_for_prompt(db_path=None):
    """LLMプロンプト用のサマリーデータを取得"""
    conn = get_connection(db_path)
    cursor = conn.cursor()

    current_fy = get_current_fiscal_year()
    prev_fy = current_fy - 1

    # 最新の報告書情報
    cursor.execute('SELECT id, report_date FROM reports ORDER BY report_date DESC LIMIT 1')
    row = cursor.fetchone()
    latest_report_id = row[0] if row else None
    report_date = row[1] if row else datetime.now().strftime('%Y-%m-%d')

    # 今年度累計売上
    cursor.execute('''
        SELECT SUM(total_sales) FROM monthly_summary
        WHERE report_id = ? AND fiscal_year = ?
    ''', (latest_report_id, current_fy))
    current_total = cursor.fetchone()[0] or 0

    # 前年度累計売上
    cursor.execute('''
        SELECT SUM(total_sales) FROM monthly_summary
        WHERE report_id = ? AND fiscal_year = ?
    ''', (latest_report_id, prev_fy))
    prev_total = cursor.fetchone()[0] or 0

    # 月別売上（今年度）
    cursor.execute('''
        SELECT month, total_sales, budget, budget_rate, yoy_rate
        FROM monthly_summary
        WHERE report_id = ? AND fiscal_year = ?
        ORDER BY CASE WHEN month >= 4 THEN month - 4 ELSE month + 8 END
    ''', (latest_report_id, current_fy))

    monthly_data = []
    for row in cursor.fetchall():
        monthly_data.append({
            'month': row[0],
            'sales': row[1] or 0,
            'budget': row[2] or 0,
            'budget_rate': row[3] or 0,
            'yoy_rate': row[4] or 0
        })

    # 属性別売上
    cursor.execute('''
        SELECT
            s.attribute,
            COUNT(DISTINCT s.id) as school_count,
            SUM(sys.total_sales) as total_sales
        FROM schools s
        LEFT JOIN school_yearly_sales sys ON sys.school_id = s.id
            AND sys.report_id = ? AND sys.fiscal_year = ?
        WHERE s.attribute IS NOT NULL AND s.attribute != ''
        GROUP BY s.attribute
        ORDER BY total_sales DESC
    ''', (latest_report_id, current_fy))

    attribute_data = []
    for row in cursor.fetchall():
        attribute_data.append({
            'attribute': row[0],
            'school_count': row[1],
            'sales': row[2] or 0
        })

    # 写真館別TOP5
    cursor.execute('''
        SELECT
            s.studio_name,
            SUM(sys.total_sales) as total_sales,
            COUNT(DISTINCT s.id) as school_count
        FROM schools s
        JOIN school_yearly_sales sys ON sys.school_id = s.id
            AND sys.report_id = ? AND sys.fiscal_year = ?
        WHERE s.studio_name IS NOT NULL
        GROUP BY s.studio_name
        ORDER BY total_sales DESC
        LIMIT 5
    ''', (latest_report_id, current_fy))

    top_studios = []
    for row in cursor.fetchall():
        top_studios.append({
            'name': row[0],
            'sales': row[1] or 0,
            'schools': row[2]
        })

    conn.close()

    return {
        'report_date': report_date,
        'fiscal_year': current_fy,
        'current_total': current_total,
        'prev_total': prev_total,
        'yoy_rate': current_total / prev_total if prev_total > 0 else 0,
        'monthly_data': monthly_data,
        'attribute_data': attribute_data,
        'top_studios': top_studios
    }


def build_analysis_prompt(summary, analytics, config):
    """分析用プロンプトを構築"""
    role = config.get('ai_consultant', {}).get('prompt_template', {}).get(
        'role',
        'あなたはスクールフォト事業の経営コンサルタントです。'
    )

    # 月別売上テキスト
    monthly_text = ""
    for m in summary['monthly_data']:
        budget_str = f"予算比{m['budget_rate']*100:.0f}%" if m['budget_rate'] else ""
        yoy_str = f"昨年比{m['yoy_rate']*100:.0f}%" if m['yoy_rate'] else ""
        monthly_text += f"  - {m['month']}月: {m['sales']:,.0f}円 {budget_str} {yoy_str}\n"

    # 属性別テキスト
    attr_text = ""
    for a in summary['attribute_data']:
        attr_text += f"  - {a['attribute']}: {a['school_count']}校, {a['sales']:,.0f}円\n"

    # 写真館TOP5テキスト
    studio_text = ""
    for i, s in enumerate(summary['top_studios'], 1):
        studio_text += f"  {i}. {s['name']}: {s['sales']:,.0f}円 ({s['schools']}校)\n"

    # 異常検知テキスト
    anomaly_text = ""
    anomalies = analytics.get('anomalies', [])[:5]
    if anomalies:
        for a in anomalies:
            change = a['change_rate'] * 100
            anomaly_text += f"  - {a['school_name']}: {a['prev_sales']:,.0f}円 → {a['current_sales']:,.0f}円 ({change:+.0f}%)\n"
    else:
        anomaly_text = "  なし\n"

    # 会員率低下テキスト
    member_text = ""
    member_trends = analytics.get('member_rate_trends', [])[:5]
    if member_trends:
        for m in member_trends:
            member_text += f"  - {m['school_name']}: {m['start_rate']*100:.1f}% → {m['current_rate']*100:.1f}%\n"
    else:
        member_text = "  なし\n"

    # 季節性分析テキスト
    seasonality = analytics.get('seasonality', {})
    season_text = ""
    for month in [4, 5, 6, 7, 8, 9, 10, 11, 12, 1, 2, 3]:
        if month in seasonality:
            s = seasonality[month]
            trend_ja = {'growing': '成長', 'stable': '安定', 'declining': '減少'}.get(s['trend'], '不明')
            season_text += f"  - {month}月: 平均{s['avg_sales']:,.0f}円 ({trend_ja}傾向)\n"

    prompt = f"""{role}

以下のスクールフォト事業の売上データを分析し、具体的なアドバイスを提供してください。

## 基本情報
- 報告書日付: {summary['report_date']}
- 対象年度: {summary['fiscal_year']}年度

## 売上概況
- 今年度累計売上: {summary['current_total']:,.0f}円
- 前年度累計売上: {summary['prev_total']:,.0f}円
- 昨年対比: {summary['yoy_rate']*100:.1f}%

## 月別売上実績
{monthly_text}

## 属性別売上（学校種別）
{attr_text}

## 写真館別売上TOP5
{studio_text}

## 注意が必要な学校（売上急減）
{anomaly_text}

## 会員率が連続低下している学校
{member_text}

## 季節性分析（過去の傾向）
{season_text}

---

上記のデータを踏まえて、以下の観点から分析とアドバイスをお願いします：

1. **現状評価**: 今年度の業績は良好か、課題があるか
2. **重点対応案件**: 優先的にフォローすべき学校や写真館
3. **季節性を踏まえた今後の見通し**: 来月以降の売上予測と対策
4. **会員率向上策**: 会員登録率を改善するための具体的施策
5. **成長機会**: 売上拡大のためのアクションプラン

回答は日本語で、具体的な数値や学校名を挙げながら実践的なアドバイスをしてください。
"""

    return prompt


def call_ollama(prompt, config):
    """Ollama APIを呼び出してレスポンスを取得"""
    ollama_config = config.get('ai_consultant', {}).get('ollama', {})
    base_url = ollama_config.get('base_url', 'http://localhost:11434')
    model = ollama_config.get('model', 'gemma2')
    timeout = ollama_config.get('timeout', 120)

    try:
        response = requests.post(
            f"{base_url}/api/generate",
            json={
                "model": model,
                "prompt": prompt,
                "stream": False,
                "options": {
                    "temperature": 0.7,
                    "num_predict": 2000
                }
            },
            timeout=timeout
        )

        if response.status_code == 200:
            result = response.json()
            return {
                'success': True,
                'content': result.get('response', ''),
                'model': model,
                'generated_at': datetime.now().isoformat()
            }
        else:
            return {
                'success': False,
                'error': f"API Error: {response.status_code}",
                'content': None
            }

    except requests.exceptions.Timeout:
        return {
            'success': False,
            'error': "タイムアウト: LLMの応答に時間がかかりすぎています",
            'content': None
        }
    except requests.exceptions.RequestException as e:
        return {
            'success': False,
            'error': f"接続エラー: {str(e)}",
            'content': None
        }


def generate_ai_advice(db_path=None):
    """AIアドバイスを生成（メイン関数）"""
    config = load_config()

    # 設定チェック
    if not config:
        return {
            'success': False,
            'error': 'config.jsonが見つかりません',
            'content': None,
            'available': False
        }

    if not config.get('ai_consultant', {}).get('enabled', False):
        return {
            'success': False,
            'error': 'AIコンサルタント機能が無効になっています',
            'content': None,
            'available': False
        }

    # Ollama接続チェック
    if not is_ollama_available(config):
        return {
            'success': False,
            'error': 'Ollamaに接続できません。Ollamaが起動しているか確認してください。',
            'content': None,
            'available': False
        }

    # データ取得
    summary = get_summary_for_prompt(db_path)
    analytics = get_all_analytics(db_path)

    # プロンプト構築
    prompt = build_analysis_prompt(summary, analytics, config)

    # LLM呼び出し
    result = call_ollama(prompt, config)
    result['available'] = True

    return result


def get_output_directory():
    """出力ディレクトリを取得（なければ作成）"""
    output_dir = Path(__file__).parent / "output"
    output_dir.mkdir(exist_ok=True)
    return output_dir


def save_as_markdown(content, fiscal_year=None):
    """Markdown形式で保存"""
    output_dir = get_output_directory()
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    fy_str = f"_{fiscal_year}年度" if fiscal_year else ""
    filename = f"AIコンサルタント分析{fy_str}_{timestamp}.md"
    filepath = output_dir / filename

    markdown_content = f"""# AIコンサルタント分析レポート

**生成日時**: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}

---

{content}

---

*このレポートはAI（Ollama）により自動生成されました。*
"""

    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(markdown_content)

    return str(filepath)


def save_as_word(content, fiscal_year=None):
    """Word形式で保存"""
    if not DOCX_AVAILABLE:
        return None

    output_dir = get_output_directory()
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    fy_str = f"_{fiscal_year}年度" if fiscal_year else ""
    filename = f"AIコンサルタント分析{fy_str}_{timestamp}.docx"
    filepath = output_dir / filename

    doc = Document()

    # タイトル
    title = doc.add_heading('AIコンサルタント分析レポート', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 生成日時
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_para.add_run(f"生成日時: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}")
    date_run.font.size = Pt(10)

    doc.add_paragraph()  # 空行

    # 本文（Markdownを簡易パース）
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('**') and line.endswith('**'):
            para = doc.add_paragraph()
            run = para.add_run(line[2:-2])
            run.bold = True
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
            doc.add_paragraph(line[3:], style='List Number')
        else:
            doc.add_paragraph(line)

    # フッター
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer_run = footer.add_run('このレポートはAI（Ollama）により自動生成されました。')
    footer_run.font.size = Pt(9)
    footer_run.italic = True

    doc.save(filepath)
    return str(filepath)


def save_as_pdf(content, fiscal_year=None):
    """PDF形式で保存（Markdownからpandocで変換）"""
    # まずMarkdownを保存
    md_path = save_as_markdown(content, fiscal_year)
    if not md_path:
        return None

    output_dir = get_output_directory()
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    fy_str = f"_{fiscal_year}年度" if fiscal_year else ""
    pdf_filename = f"AIコンサルタント分析{fy_str}_{timestamp}.pdf"
    pdf_path = output_dir / pdf_filename

    # pandocでPDF変換を試みる
    try:
        result = subprocess.run(
            ['pandoc', md_path, '-o', str(pdf_path),
             '--pdf-engine=xelatex',
             '-V', 'mainfont=Noto Sans CJK JP',
             '-V', 'geometry:margin=2cm'],
            capture_output=True,
            text=True,
            timeout=60
        )
        if result.returncode == 0 and pdf_path.exists():
            return str(pdf_path)
    except (subprocess.TimeoutExpired, FileNotFoundError):
        pass

    # pandocが失敗した場合はMarkdownパスを返す
    return None


def export_ai_advice(format='markdown', db_path=None):
    """AIアドバイスを生成して指定形式でエクスポート

    Args:
        format: 'markdown', 'word', 'pdf' のいずれか
        db_path: データベースパス

    Returns:
        dict: {
            'success': bool,
            'filepath': str or None,
            'format': str,
            'error': str or None
        }
    """
    # AIアドバイス生成
    result = generate_ai_advice(db_path)

    if not result['success']:
        return {
            'success': False,
            'filepath': None,
            'format': format,
            'error': result.get('error', '不明なエラー')
        }

    content = result['content']
    config = load_config()
    fiscal_year = get_current_fiscal_year()

    # 形式に応じて保存
    filepath = None
    actual_format = format

    if format == 'word':
        if DOCX_AVAILABLE:
            filepath = save_as_word(content, fiscal_year)
        else:
            # python-docxがない場合はMarkdownにフォールバック
            filepath = save_as_markdown(content, fiscal_year)
            actual_format = 'markdown'
    elif format == 'pdf':
        filepath = save_as_pdf(content, fiscal_year)
        if not filepath:
            # PDF変換失敗時はMarkdownにフォールバック
            filepath = save_as_markdown(content, fiscal_year)
            actual_format = 'markdown'
    else:  # markdown
        filepath = save_as_markdown(content, fiscal_year)

    return {
        'success': True,
        'filepath': filepath,
        'format': actual_format,
        'error': None,
        'model': result.get('model'),
        'generated_at': result.get('generated_at')
    }


if __name__ == '__main__':
    import sys

    print("AIコンサルタント分析を実行中...")

    # コマンドライン引数で形式を指定可能
    format_arg = 'markdown'
    if len(sys.argv) > 1:
        if sys.argv[1] in ['word', 'pdf', 'markdown', 'md']:
            format_arg = 'word' if sys.argv[1] == 'word' else \
                        'pdf' if sys.argv[1] == 'pdf' else 'markdown'

    result = export_ai_advice(format=format_arg)

    if result['success']:
        print(f"\n✅ 分析レポートを出力しました")
        print(f"   形式: {result['format']}")
        print(f"   ファイル: {result['filepath']}")
        print(f"   モデル: {result['model']}")
        print(f"   生成日時: {result['generated_at']}")
    else:
        print(f"❌ エラー: {result['error']}")
